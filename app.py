import os
from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory, g
from docx import Document
import sqlite3, uuid
from cololog import cololog
from datetime import datetime

logger = cololog(__name__, path_print=False, log_dir='logs', log_to_file=True)

app = Flask(__name__)

# Секретный ключ для работы с сессиями
app.secret_key = os.urandom(24)

# Папка для загрузки файлов
UPLOAD_FOLDER = 'uploads'

# Папка для загрузки файлов
EXPENSES_FOLDER = 'expenses'


if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
    logger.info(f"Папка для загрузки файлов '{UPLOAD_FOLDER}' создана.")

if not os.path.exists(EXPENSES_FOLDER):
    os.makedirs(EXPENSES_FOLDER)
    logger.info(f"Папка для расходов '{EXPENSES_FOLDER}' создана.")


# Подключение к базе данных
def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect('database.db')
        g.db.row_factory = sqlite3.Row  # Для удобства работы с результатами запроса (словари)
        logger.info("Подключение к базе данных установлено.")
    return g.db

# Закрытие подключения к базе данных
@app.teardown_appcontext
def close_db(error):
    db = getattr(g, 'db', None)
    if db is not None:
        db.close()
        logger.info("Подключение к базе данных закрыто.")


def init_db():
    with app.app_context():
        db = get_db()
        with app.open_resource('sqlite/transactions.sql', mode='r') as f:
            db.cursor().executescript(f.read())
        db.commit()
        logger.info("База данных инициализирована.")

app.secret_key = 'VOvWd8xuPf'

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['EXPENSES_FOLDER'] = EXPENSES_FOLDER


# Пароль для входа
ADMIN_PASSWORD = 'h5wwLy'  # Лучше хранить в переменных окружения

# Функция для создания папки с номером заявки
def create_directory_for_order(order_number):
    order_dir = os.path.join(app.config['UPLOAD_FOLDER'], order_number)
    if not os.path.exists(order_dir):
        os.makedirs(order_dir)
        logger.info(f"Создана папка для заказа с номером {order_number}.")
    return order_dir


# Проверка аутентификации
def is_authenticated():
    return 'authenticated' in session and session['authenticated']


# Главная страница
@app.route('/')
def index():
    orders = os.listdir(UPLOAD_FOLDER)
    logger.info(f"Загружены все заказы: {orders}")
    return render_template('index.html', orders=orders)


# Страница входа (admin)
@app.route('/admin', methods=['GET', 'POST'])
def admin():
    if is_authenticated():
        logger.info("Пользователь уже авторизован, редирект на главную страницу.")
        return redirect(url_for('main'))  # Переход на страницу загрузки при авторизации

    if request.method == 'POST':
        password = request.form.get('password')
        if password == ADMIN_PASSWORD:
            session['authenticated'] = True
            logger.info("Пользователь авторизовался.")
            return redirect(url_for('main'))  # Переход на страницу загрузки при успешном входе
        else:
            logger.warning("Неверный пароль при попытке входа.")
            return "Неверный пароль. Попробуйте снова."

    return render_template('admin_login.html')


# Функция для создания папки с уникальным ID траты
def create_directory_for_expense(expense_id):
    expense_dir = os.path.join(EXPENSES_FOLDER, expense_id)
    if not os.path.exists(expense_dir):
        os.makedirs(expense_dir)
        logger.info(f"Создана папка для траты с ID {expense_id}.")
    return expense_dir


@app.route('/expenses/<expense_id>/<filename>')
def serve_expense_file(expense_id, filename):
    # Указываем путь к файлам в папке 'expenses'
    file_path = os.path.join('expenses', expense_id, filename)
    
    # Проверяем, существует ли файл
    if not os.path.exists(file_path):
        logger.error(f"Файл {filename} не найден в папке траты с ID {expense_id}.")
        return "Файл не найден", 404
    
    # Возвращаем файл
    return send_from_directory(os.path.dirname(file_path), os.path.basename(file_path))


@app.route('/expense/<expense_id>', methods=['GET'])
def view_expense(expense_id):
    if not is_authenticated():
        logger.warning("Попытка просмотра расходов без авторизации.")
        return redirect(url_for('admin'))  # Если не авторизован, перенаправление на страницу логина
    
    # Путь к папке с заказом в директории 'expenses'
    order_dir = os.path.join('expenses', expense_id)  # Папка с заказом внутри expenses
    if not os.path.exists(order_dir):
        logger.error(f"Трата с ID {expense_id} не найдена.")
        return "Трата не найден", 404

    # Путь к DOCX документу
    doc_filename = f'{expense_id}_expense_document.docx'
    doc_path = os.path.join(order_dir, doc_filename)
    
    if not os.path.exists(doc_path):
        logger.error(f"Документ с ID {expense_id} не найден.")
        return "Документ не найден", 404

    # Открытие и чтение документа
    doc = Document(doc_path)

    order_data = {}
    order_data['services'] = []

    for para in doc.paragraphs:
        text = para.text.strip()  # Убираем лишние пробелы

        order_data['uuid'] = expense_id
        if text.startswith('ФИО:'):
            order_data['name'] = text.replace('ФИО:', '').strip()
        elif text.startswith('Описание:'):
            order_data['description'] = text.replace('Описание:', '').strip()
        elif text.startswith('Сумма:'):
            order_data['total_cost'] = text.replace('Сумма:', '').strip()
        elif text.startswith('Чек:'):
            order_data['receipt_filename'] = text.replace('Чек:', '').strip()
        elif text.startswith('Дата:'):
            order_data['date'] = text.replace('Дата:', '').strip()

    logger.info(f"Данные по трате {expense_id}: {order_data}")

    print(order_data['uuid'])

    # Формируем данные для передачи в шаблон
    expense_data = {
        'uuid': expense_id,
        'name': order_data['name'],
        'description': order_data['description'],
        'amount': order_data['total_cost'],
        'transaction_date': order_data['date']
    }

    return render_template('view_expense.html', expense_data=expense_data, filename=order_data["receipt_filename"])


@app.route('/main', methods=['GET', 'POST'])
def main():
    if not is_authenticated():
        logger.warning("Попытка перейти на главную без авторизации.")
        return redirect(url_for('admin'))
    
    # Получаем запрос из формы поиска
    search_query = request.args.get('search', '').lower()

    # Выполняем запросы для подсчета доходов и расходов
    db = get_db()

    # Подсчет всех доходов
    total_income = db.execute('''
        SELECT SUM(amount) FROM transactions WHERE transaction_type = 'income'
    ''').fetchone()[0] or 0  # Если результата нет, то сумма 0

    # Подсчет всех расходов
    total_expenses = db.execute('''
        SELECT SUM(amount) FROM transactions WHERE transaction_type = 'expense'
    ''').fetchone()[0] or 0  # Если результата нет, то сумма 0

    # Итог = доходы - расходы
    total_balance = total_income - total_expenses

    # Получаем список всех заявок
    orders = os.listdir(UPLOAD_FOLDER)
    filtered_orders = []

    # Получаем список всех трат
    expenses = os.listdir(EXPENSES_FOLDER)
    filtered_expenses = []

    # Фильтруем заявки
    for order in orders:
        order_dir = os.path.join(app.config['UPLOAD_FOLDER'], order)
        
        if os.path.isdir(order_dir):
            doc_file_path = None
            for file_ in os.listdir(order_dir):
                if file_.endswith('.docx'):
                    doc_file_path = os.path.join(order_dir, file_)
                    break
            
            if doc_file_path:
                doc = Document(doc_file_path)
                text_content = ' '.join([para.text.lower() for para in doc.paragraphs])  # Собираем все тексты из документа в одну строку
                
                if search_query in text_content:
                    filtered_orders.append(order)

    # Фильтруем траты
    for expense in expenses:
        expense_dir = os.path.join(app.config['EXPENSES_FOLDER'], expense)
        
        if os.path.isdir(expense_dir):
            doc_file_path = None
            for file_ in os.listdir(expense_dir):
                if file_.endswith('.docx'):
                    doc_file_path = os.path.join(expense_dir, file_)
                    break
            
            if doc_file_path:
                doc = Document(doc_file_path)
                text_content = ' '.join([para.text.lower() for para in doc.paragraphs])  # Собираем все тексты из документа в одну строку
                
                if search_query in text_content:
                    filtered_expenses.append(expense)

    # Если не было найдено совпадений по запросу, показываем все заявки и все траты
    if search_query:
        orders = filtered_orders
        expenses = filtered_expenses

    logger.info(f"Итоги: Доходы: {total_income}, Расходы: {total_expenses}, Баланс: {total_balance}")

    # Возвращаем шаблон с отфильтрованными заявками и тратами
    return render_template('admin_index.html', 
                           orders=orders, 
                           expenses=expenses,
                           total_income=total_income,
                           total_expenses=total_expenses,
                           total_balance=total_balance)


# Страница загрузки заявки (upload)
@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if not is_authenticated():
        logger.warning("Попытка загрузить заявку без авторизации.")
        return redirect(url_for('admin'))  # Если не авторизован, перенаправление на страницу логина

    if request.method == 'POST':
        # Получение данных из формы
        name = request.form['name']
        name_customer = request.form['name_customer']
        address = request.form['address']
        services = request.form.getlist('services[]')
        prices = request.form.getlist('prices[]')
        quantities = request.form.getlist('quantities[]')
        sums = request.form.getlist('sums[]')
        warranties = request.form.getlist('warranties[]')
        discount = float(request.form['discount'])
        total_cost = float(request.form['total_cost'])
        discounted_cost = float(request.form['discounted_cost'])
        recommendations = request.form['recommendations']
        order_number = request.form['order_number']
        order_date_str = request.form['order_date']  # Дата заявки

        # Преобразуем строку с датой в объект datetime
        try:
            order_date = datetime.strptime(order_date_str, "%d.%m.%Y")
        except ValueError:
            logger.error(f"Ошибка формата даты для заявки {order_number}. Ожидаемый формат: ДД.ММ.ГГГГ.")
            return "Ошибка: Неверный формат даты. Используйте формат: ДД.ММ.ГГГГ"

        # Получение файлов
        receipt_file = request.files['receipt']
        document_file1 = request.files['document1']
        document_file2 = request.files['document2']

        # Создание папки для заявок
        order_dir = create_directory_for_order(order_number)

        # Генерация новых имен файлов на английском языке
        receipt_filename = 'receipt.jpg'
        document_filename1 = 'document1.jpg'
        document_filename2 = 'document2.jpg'

        # Путь к сохранению файлов
        receipt_path = os.path.join(order_dir, receipt_filename)
        document_path1 = os.path.join(order_dir, document_filename1)
        document_path2 = os.path.join(order_dir, document_filename2)

        # Сохранение файлов
        receipt_file.save(receipt_path)
        document_file1.save(document_path1)
        document_file2.save(document_path2)

        # Генерация DOCX документа
        doc = Document()
        doc.add_heading('Документ о выполненных работах', 0)
        doc.add_paragraph(f'Исполнитель: {name}')
        doc.add_paragraph(f'Заказчик: {name_customer}')
        doc.add_paragraph(f'Адрес выполнения работ: {address}')
        doc.add_paragraph(f'Рекомендации: {recommendations}')
        doc.add_paragraph(f'Номер заявки: {order_number}')
        doc.add_paragraph(f'Дата выполнения работ: {order_date.strftime("%Y-%m-%d")}')

        # Добавляем информацию о услугах
        for i in range(len(services)):
            doc.add_paragraph(f'Услуга: {services[i]}')
            doc.add_paragraph(f'Цена: {prices[i]}')
            doc.add_paragraph(f'Количество: {quantities[i]}')
            doc.add_paragraph(f'Сумма: {sums[i]}')
            doc.add_paragraph(f'Гарантия: {warranties[i]}')

        # Добавляем стоимость услуг и скидку
        doc.add_paragraph(f'Стоимость услуг: {total_cost}')
        doc.add_paragraph(f'Скидка: {discount}')
        doc.add_paragraph(f'К оплате: {discounted_cost}')

        # Вставляем ссылки на файлы
        doc.add_paragraph(f'Чек: {receipt_filename}')
        doc.add_paragraph(f'Документ (сторона 1): {document_filename1}')
        doc.add_paragraph(f'Документ (сторона 2): {document_filename2}')

        # Сохранение документа
        doc_filename = f'{order_number}_document.docx'
        doc_path = os.path.join(order_dir, doc_filename)
        doc.save(doc_path)

        # Сохранение данных в базу данных
        db = get_db()
        # Добавление данных о доходах/расходах в отдельную таблицу
        db.execute('''
            INSERT INTO transactions (transaction_type, amount, receipt_filename, transaction_date, uuid)
            VALUES (?, ?, ?, ?, ?)
        ''', ("income", discounted_cost, receipt_filename, datetime.now().strftime('%Y-%m-%d'), order_number))
        db.commit()

        logger.info(f"Заявка с номером {order_number} успешно загружена.")

        return redirect(url_for('main'))

    return render_template('upload.html')


# Добавляем маршрут для обработки статических файлов
@app.route('/uploads/<order_number>/<filename>')
def uploaded_file(order_number, filename):
    return send_from_directory(os.path.join(app.config['UPLOAD_FOLDER'], order_number), filename)


@app.route('/view_order/<order_number>', methods=['GET'])
def view_order(order_number):
    if not is_authenticated():
        return redirect(url_for('admin'))
    
    # Путь к папке с заказом в директории 'uploads'
    order_dir = os.path.join('uploads', order_number)  # Папка с заказом внутри uploads
    if not os.path.exists(order_dir):
        return "Заказ не найден", 404

    # Путь к DOCX документу
    doc_filename = f'{order_number}_document.docx'
    doc_path = os.path.join(order_dir, doc_filename)
    
    if not os.path.exists(doc_path):
        return "Документ не найден", 404

    # Открытие и чтение документа
    doc = Document(doc_path)

    order_data = {}
    order_data['services'] = []

    for para in doc.paragraphs:
        text = para.text.strip()  # Убираем лишние пробелы

        if text.startswith('Исполнитель:'):
            order_data['name'] = text.replace('Исполнитель:', '').strip()
        elif text.startswith('Заказчик:'):
            order_data['name_customer'] = text.replace('Заказчик:', '').strip()
        elif text.startswith('Адрес выполнения работ:'):
            order_data['address'] = text.replace('Адрес выполнения работ:', '').strip()
        elif text.startswith('Рекомендации:'):
            order_data['recommendations'] = text.replace('Рекомендации:', '').strip()
        elif text.startswith('Номер заявки:'):
            order_data['order_number'] = text.replace('Номер заявки:', '').strip()
        elif text.startswith('Дата выполнения работ:'):
            order_data['order_date'] = text.replace('Дата выполнения работ:', '').strip()
        elif text.startswith('Стоимость услуг:'):
            order_data['total_cost'] = text.replace('Стоимость услуг:', '').strip()
        elif text.startswith('Скидка:'):
            order_data['discount'] = text.replace('Скидка:', '').strip()
        elif text.startswith('К оплате:'):
            order_data['discounted_cost'] = text.replace('К оплате:', '').strip()
        elif text.startswith('Чек:'):
            order_data['receipt_filename'] = text.replace('Чек:', '').strip()
        elif text.startswith('Документ (сторона 1):'):
            order_data['document_filename1'] = text.replace('Документ (сторона 1):', '').strip()
        elif text.startswith('Документ (сторона 2):'):
            order_data['document_filename2'] = text.replace('Документ (сторона 2):', '').strip()
        elif text.startswith('Услуга:'):
            service = {}
            service['name'] = text.replace('Услуга:', '').strip()
            order_data['services'].append(service)
        elif text.startswith('Цена:') and order_data['services']:
            order_data['services'][-1]['price'] = text.replace('Цена:', '').strip()
        elif text.startswith('Количество:') and order_data['services']:
            order_data['services'][-1]['quantity'] = text.replace('Количество:', '').strip()
        elif text.startswith('Сумма:') and order_data['services']:
            order_data['services'][-1]['sum'] = text.replace('Сумма:', '').strip()
        elif text.startswith('Гарантия:') and order_data['services']:
            order_data['services'][-1]['warranty'] = text.replace('Гарантия:', '').strip()
        

    # Формируем URL-адреса для файлов
    receipt_url = url_for('static', filename=f'uploads/{order_number}/{order_data["receipt_filename"]}')
    document_url1 = url_for('static', filename=f'uploads/{order_number}/{order_data["document_filename1"]}')
    document_url2 = url_for('static', filename=f'uploads/{order_number}/{order_data["document_filename2"]}')

    # Возвращаем информацию на страницу
    return render_template('view_order.html',
            order_data=order_data, 
            receipt_url=receipt_url,
            document_url1=document_url1,
            document_url2=document_url2
    )


# Страница для добавления расходов
@app.route('/add_expense', methods=['GET', 'POST'])
def add_expense():
    if not is_authenticated():
        return redirect(url_for('admin'))  # Если не авторизован, перенаправление на страницу логина

    if request.method == 'POST':
        # Получение данных из формы
        name = request.form['name']
        description = request.form['description']
        amount = float(request.form['amount'])

        # Генерация уникального ID для траты
        expense_id = str(uuid.uuid4())  # Генерация уникального ID для траты

        # Получение файла чека
        receipt_file = request.files['receipt']
        
        # Создание папки для траты
        expense_dir = create_directory_for_expense(expense_id)

        # Сохранение чека
        receipt_filename = 'receipt.jpg'
        receipt_path = os.path.join(expense_dir, receipt_filename)
        receipt_file.save(receipt_path)

        # Генерация документа (если необходимо) или сохранение данных
        doc = Document()
        doc.add_heading('Документ о расходах', 0)
        doc.add_paragraph(f'ФИО: {name}')
        doc.add_paragraph(f'Описание: {description}')
        doc.add_paragraph(f'Сумма: {amount}')
        doc.add_paragraph(f'Чек: {receipt_filename}')
        current_datetime = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
        doc.add_paragraph(f'Дата: {current_datetime}')

        # Генерация имени файла для документа
        doc_filename = f'{expense_id}_expense_document.docx'
        doc_path = os.path.join(expense_dir, doc_filename)
        doc.save(doc_path)

        # Сохранение данных в базу данных
        db = get_db()
        # Добавление данных о доходах/расходах в отдельную таблицу
        db.execute('''
            INSERT INTO transactions (transaction_type, amount, receipt_filename, transaction_date, uuid)
            VALUES (?, ?, ?, ?, ?)
        ''', ("expense", amount, receipt_filename, datetime.now().strftime('%Y-%m-%d'), expense_id))
        db.commit()

        return redirect(url_for('admin'))  # Возврат на главную страницу

    return render_template('add_expense.html')


# Выход из системы
@app.route('/logout')
def logout():
    logger.info("Пользователь вышел из системы.")
    session.pop('authenticated', None)
    return redirect(url_for('index'))


if __name__ == '__main__':
    init_db()
    try:
        app.run(host='0.0.0.0', port=443, ssl_context=('/etc/letsencrypt/live/repair-31.ru/fullchain.pem', '/etc/letsencrypt/live/repair-31.ru/privkey.pem'))
    except:
        app.run(debug=True)